# -*- coding: utf-8 -*-
"""
Script chuyển đổi Markdown (.md) sang Word (.docx) chuyên nghiệp
sử dụng python-docx, định dạng font Times New Roman, cỡ chữ, lề, giãn dòng chuẩn Việt Nam.
"""
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

MD_PATH = Path("docs/report/Bao_Cao_Phan_Tich_Du_Lieu_Nhom2.md")
DOCX_PATH = Path("docs/report/Bao_Cao_Phan_Tich_Du_Lieu_Nhom2.docx")

def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho ô bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Đặt padding/margin cho ô bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Thiết lập viền mảnh màu xám cho bảng."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_run(run, font_name="Times New Roman", size_pt=12, color_rgb=None, bold=False, italic=False):
    """Định dạng phông chữ cho Run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color_rgb:
        run.font.color.rgb = color_rgb

def parse_markdown_to_docx():
    print(f"Reading Markdown file from: {MD_PATH.resolve()}...")
    if not MD_PATH.exists():
        print(f"Error: {MD_PATH} does not exist!")
        return

    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Thiết lập lề trang A4 chuẩn (Top/Bottom 2.5cm, Left 3.0cm, Right 2.0cm)
    for section in doc.sections:
        section.top_margin = Inches(0.98)     # ~2.5 cm
        section.bottom_margin = Inches(0.98)  # ~2.5 cm
        section.left_margin = Inches(1.18)    # ~3.0 cm
        section.right_margin = Inches(0.79)   # ~2.0 cm
        
    lines = content.split('\n')
    
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 1. Bỏ qua các dòng gạch phân cách
        if line == '---' or line == '***':
            i += 1
            continue
            
        # 2. Xử lý Bảng biểu (Markdown Table)
        if line.startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # Kết thúc bảng, vẽ bảng vào Word
            in_table = False
            draw_word_table(doc, table_rows)
            table_rows = []
            
        if not line:
            i += 1
            continue
            
        # 3. Xử lý Heading
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            title_text = heading_match.group(2)
            title_text = clean_inline_markdown(title_text)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title_text)
            if level == 1:
                format_run(run, size_pt=18, bold=True, color_rgb=RGBColor(26, 54, 93))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if "BÁO CÁO PHÂN TÍCH" in title_text.upper() else WD_ALIGN_PARAGRAPH.LEFT
            elif level == 2:
                format_run(run, size_pt=14, bold=True, color_rgb=RGBColor(44, 82, 130))
            else:
                format_run(run, size_pt=12, bold=True, color_rgb=RGBColor(74, 85, 104))
                
            i += 1
            continue
            
        # 4. Xử lý Ảnh chèn (Markdown Image)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
        if img_match:
            caption = img_match.group(1)
            img_path_str = img_match.group(2)
            # Normalize path
            img_path = Path("docs/report") / img_path_str
            if img_path.exists():
                print(f"Inserting image: {img_path.name}...")
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(4)
                
                # Resize ảnh cho cân đối trên trang
                p_img.add_run().add_picture(str(img_path), width=Inches(5.8))
                
                # Caption của ảnh
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(caption)
                format_run(run_cap, size_pt=10, italic=True, color_rgb=RGBColor(100, 100, 100))
            else:
                print(f"Warning: Image path not found: {img_path}")
            i += 1
            continue
            
        # 5. Xử lý Danh sách không thứ tự (Bullet List)
        list_match = re.match(r'^[-*+]\s+(.*)$', line)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            add_formatted_text(p, text)
            i += 1
            continue
            
        # 6. Xử lý Paragraph thông thường
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_formatted_text(p, line)
        
        i += 1

    # Nếu file kết thúc mà còn bảng dở dang
    if table_rows:
        draw_word_table(doc, table_rows)
        
    doc.save(DOCX_PATH)
    print(f"Successfully exported report to Word at: {DOCX_PATH.resolve()}")

def clean_inline_markdown(text):
    """Loại bỏ các tag in đậm nghiêng markdown để hiển thị text thuần túy."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def add_formatted_text(paragraph, md_text):
    """Phân tích văn bản nội dòng (Bold, Italic, Code, Link) và thêm các Run tương ứng."""
    # Tokenize các pattern: bold (**...**), italic (*...*), code (`...`)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', md_text)
    
    for token in tokens:
        if not token:
            continue
        
        # Bold
        if token.startswith('**') and token.endswith('**'):
            t = token[2:-2]
            run = paragraph.add_run(t)
            format_run(run, bold=True)
        # Italic
        elif token.startswith('*') and token.endswith('*'):
            t = token[1:-1]
            run = paragraph.add_run(t)
            format_run(run, italic=True)
        # Inline code
        elif token.startswith('`') and token.endswith('`'):
            t = token[1:-1]
            run = paragraph.add_run(t)
            format_run(run, font_name="Courier New", size_pt=10.5, color_rgb=RGBColor(150, 40, 40))
        # Markdown Link
        elif token.startswith('[') and '(' in token:
            link_match = re.match(r'^\[(.*?)\]\((.*?)\)$', token)
            if link_match:
                t = link_match.group(1)
                run = paragraph.add_run(t)
                format_run(run, color_rgb=RGBColor(44, 82, 130), bold=True)
            else:
                run = paragraph.add_run(token)
                format_run(run)
        # Normal text
        else:
            run = paragraph.add_run(token)
            format_run(run)

def draw_word_table(doc, rows):
    """Vẽ bảng Markdown vào file Word với định dạng chuyên nghiệp."""
    # Lọc bỏ dòng phân tách tiêu đề |---|---|
    data_rows = []
    for r in rows:
        if re.match(r'^\|[\s:-|]+$', r):
            continue
        data_rows.append(r)
        
    if not data_rows:
        return
        
    # Phân tích số cột
    first_row_cols = [c.strip() for c in data_rows[0].split('|')[1:-1]]
    num_cols = len(first_row_cols)
    num_rows = len(data_rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Điền dữ liệu
    for r_idx, row_str in enumerate(data_rows):
        cols = [c.strip() for c in row_str.split('|')[1:-1]]
        # Điền các ô còn thiếu nếu có lỗi căn chỉnh markdown
        while len(cols) < num_cols:
            cols.append('')
            
        row = table.rows[r_idx]
        
        # Độ cao dòng
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '360')  # Chiều cao tối thiểu
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        is_header = (r_idx == 0)
        
        for c_idx, val in enumerate(cols[:num_cols]):
            cell = row.cells[c_idx]
            cell.text = ""  # Clear default text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            # Format text in cell
            clean_val = clean_inline_markdown(val)
            run = p.add_run(clean_val)
            
            if is_header:
                set_cell_background(cell, "1A365D")  # Dark blue header
                format_run(run, font_name="Times New Roman", size_pt=10.5, color_rgb=RGBColor(255, 255, 255), bold=True)
            else:
                # Xen kẽ màu dòng (zebra striping)
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F7FAFC")
                else:
                    set_cell_background(cell, "FFFFFF")
                format_run(run, font_name="Times New Roman", size_pt=10, color_rgb=RGBColor(45, 55, 72))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

if __name__ == "__main__":
    parse_markdown_to_docx()
